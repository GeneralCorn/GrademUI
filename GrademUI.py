import streamlit as st
from io import BytesIO
import random
import decimal
import time
import collections
import pandas as pd
import docx
from streamlit_echarts import st_echarts
# Favicon and Headings
st.set_page_config(page_title="Gradem", page_icon="💎", layout="wide")

# Headings
st.title("Welcome to Gradem!")
st.header("Gradem is a centralized IB MYP design comments generator")

with st.sidebar.expander("Please download the template files"):
    with open("./Templates.zip", "rb") as file:
        btn = st.download_button(
            label="Download",
            data=file,
            file_name="Templates.zip",
            mime="file/zip"
        )
st.sidebar.caption("Please name sheets in excel document by class, i.e. (G7.1)")

st.caption("Scroll to the very bottom for some visualizations")

# self.Basic Values
periodInput = st.sidebar.selectbox("Select Period: ", ("semester", "year"))

options = st.sidebar.radio(
    "Specify unit(s) or use time period in question?",
    ("Specify Unit(s)", "Use Time Period"))

multipleunits = False
if options == "Specify Unit(s)":
    unitinput = st.sidebar.text_input("Input Subject Unit:")
    if unitinput == "":
        st.warning("Please Input Subject Unit")
        st.stop()
    multipleunits = st.sidebar.checkbox("Multiple Units?")

st.sidebar.caption("i.e. Landscapes, Web Design")

# File I/O
stu = st.sidebar.file_uploader("Select student file", type=["csv", "xlsx"])
st.sidebar.caption("In the form of student.csv or student.xlsx")

sentences = st.sidebar.file_uploader(
    "Select commentbank file", type=["csv", "xlsx"])
st.sidebar.caption("In the form of sentences.csv or sentences.xlsx")

if "studentfile" not in st.session_state:
    st.session_state["studentfile"] = stu

# Convert input files into list and reformat accordingly
if stu is not None:
    if stu.name[-4:] == ".csv":
        df = pd.read_csv(stu)
        studentinfo = df.values.tolist()
    else:
        #persistent variables
        studentinfo = []
        sheets = pd.ExcelFile(stu)
        sheetnames = sheets.sheet_names
        es = "Export Summary"
        if es in sheetnames:
            sheetnames.remove(es)
        collectiveInfo = st.selectbox("Select Class:",(sheetnames))
        df = pd.read_excel(stu, collectiveInfo)
        studentinfo.extend(df.values.tolist())

else:
    st.warning("Upload student file")
    st.stop()


if sentences is not None:
    if sentences.name[-4:] == ".csv":
        df1 = pd.read_csv(sentences)
        commentbank = df1.values.tolist()
        for i in commentbank:
            del i[0]
    else:
        df1 = pd.read_excel(sentences)
        commentbank = df1.values.tolist()
        for i in commentbank:
            del i[0]
else:
    st.warning("Upload commentbank file")
    st.stop()

#<-------->
exportComments = docx.Document()
exportComments.add_heading(collectiveInfo)
for key in st.session_state:
    if key not in('studentfile','load'):
        exportComments.add_heading(key, level=2)
        paragraph = exportComments.add_paragraph(str(st.session_state[key]))
        paragraph.alignment = 4

upPeriod = periodInput.capitalize()
target_stream = BytesIO()
exportComments.save(target_stream)

st.download_button(
    "Export as Word file",
    target_stream,
    mime="application/msword",
    file_name=f"{collectiveInfo} MYP Design {upPeriod} Comments.docx",
    help="Note, will regenerate comments")

# student class, each object has unique set of info list based on object

class student:

    def __init__(self, col):
        self.col = col
        self.intlist = []
        for i in range(2, 6):
            self.intlist.append(studentinfo[col][i])

        self.A = self.intlist[0]
        self.B = self.intlist[1]
        self.C = self.intlist[2]
        self.D = self.intlist[3]
        self.fn = studentinfo[col][0]
        self.ln = studentinfo[col][1]

    def totalMarks(self):
        return sum(self.intlist)

    def finalGrade(self):
        x = sum(self.intlist)
        if x <= 4:
            return 1
        elif x <= 8:
            return 2
        elif x <= 13:
            return 3
        elif x <= 17:
            return 4
        elif x <= 22:
            return 5
        elif x <= 26:
            return 6
        elif x <= 32:
            return 7

    def deviation(self):
        deviation = max(self.intlist) - min(self.intlist)
        return deviation

    def fs(self):
        st1 = ""
        if self.finalGrade() == 7:
            st1 = commentbank[0][random.randint(0, 4)]
        elif self.finalGrade() == 6:
            st1 = commentbank[1][random.randint(0, 4)]
        elif self.finalGrade() == 5:
            st1 = commentbank[2][random.randint(0, 4)]
        elif self.finalGrade() == 4:
            st1 = commentbank[3][random.randint(0, 4)]
        elif self.finalGrade() == 3:
            st1 = commentbank[4][random.randint(0, 4)]
        elif self.finalGrade() == 2:
            st1 = commentbank[5][random.randint(0, 4)]
        else:
            st1 = commentbank[6][random.randint(0, 4)]
        return st1

    def ss(self):
        st2 = ""
        tg = int(studentinfo[self.col][6])
        if self.finalGrade() > tg:
            st2 = commentbank[7][0]
        elif self.finalGrade() == tg:
            st2 = commentbank[7][1]
        else:
            st2 = commentbank[7][2]
        return st2

    def ts(self):
        eff = int(studentinfo[self.col][12])
        st3 = ""
        if eff == 1:
            if self.deviation() <= 1:
                st3 = commentbank[10][0]
            elif self.deviation() == 2:
                st3 = commentbank[10][1]
            else:
                st3 = commentbank[10][2]
        elif eff == 2:
            if self.deviation() <= 2:
                st3 = commentbank[9][0]
            elif self.deviation() == 2:
                st3 = commentbank[9][1]
            else:
                st3 = commentbank[9][2]
        elif eff == 3:
            if self.deviation() <= 1:
                st3 = commentbank[8][0]
            elif self.deviation() == 2:
                st3 = commentbank[8][1]
            else:
                st3 = commentbank[8][2]
        return st3

    def fos(self):
        st4 = ""
        maxGrade = max(self.intlist)
        hci = self.intlist.index(maxGrade)

        if self.finalGrade() in range(6, 8):
            if (self.A == self.B & self.B == self.C & self.C == self.D):
                st4 = commentbank[11][0]
            else:
                st4 = commentbank[11][hci]
        elif self.finalGrade() in range(4, 6):
            if (self.A == self.B & self.B == self.C & self.C == self.D):
                st4 = commentbank[12][0]
            else:
                st4 = commentbank[12][hci]
        else:
            if (self.A == self.B & self.B == self.C & self.C == self.D):
                st4 = commentbank[13][0]
            else:
                st4 = commentbank[13][hci]
        return st4

    def fis(self):
        st5 = ""
        minGrade = min(self.intlist)
        lci = self.intlist.index(minGrade)

        if self.finalGrade() in range(6, 8):
            if self.A == self.B & self.B == self.C & self.C == self.D:
                st5 = commentbank[14][0]
            else:
                st5 = commentbank[14][lci]
        elif self.finalGrade() in range(4, 6):
            if self.A == self.B & self.B == self.C & self.C == self.D:
                st5 = commentbank[15][0]
            else:
                st5 = commentbank[15][lci]
        else:
            if (self.A == self.B & self.B == self.C & self.C == self.D):
                st5 = commentbank[16][0]
            else:
                st5 = commentbank[16][lci]
        return st5

    def sis(self):
        st6 = ""
        if studentinfo[self.col][10] == "EE":
            st6 = commentbank[17][0]
        elif studentinfo[self.col][10] == "ME":
            st6 = commentbank[17][1]
        elif studentinfo[self.col][10] == "AE":
            st6 = commentbank[17][2]
        elif studentinfo[self.col][10] == "BE":
            st6 = commentbank[17][3]
        return st6

    def ses(self):
        st7 = ""
        if studentinfo[self.col][11] == "EE":
            st7 = commentbank[18][0]
        elif studentinfo[self.col][11] == "ME":
            st7 = commentbank[18][1]
        elif studentinfo[self.col][11] == "AE":
            st7 = commentbank[18][2]
        elif studentinfo[self.col][11] == "BE":
            st7 = commentbank[18][3]
        return st7

    def final(self):
        st8 = ""
        if self.finalGrade() == 7:
            st8 = commentbank[19][random.randint(0, 4)]
        elif self.finalGrade() == 6:
            st8 == commentbank[20][random.randint(0, 4)]
        elif self.finalGrade() == 5:
            st8 == commentbank[21][random.randint(0, 4)]
        elif self.finalGrade() == 4:
            st8 == commentbank[22][random.randint(0, 4)]
        elif self.finalGrade() == 3:
            st8 == commentbank[23][random.randint(0, 4)]
        elif self.finalGrade() == 2:
            st8 == commentbank[24][random.randint(0, 4)]
        else:
            st8 == commentbank[25][random.randint(0, 4)]
        return st8

    def finalComment(self):
        cumlt = self.fs() + self.ss() + self.ts() + self.fos() + \
            self.fis() + self.sis() + self.ses() + self.final()
        named = cumlt.replace("Student!", studentinfo[self.col][0])
        atl1 = named.replace("ATL!", studentinfo[self.col][8])
        atl2 = atl1.replace("ATL2!", studentinfo[self.col][9])
        
        if multipleunits:
            temp = atl2.replace("Unit!", unitinput) 
            unit = temp.replace("unit", "units")
        elif options == "Use Time Period":
            unit = atl2.replace("Unit! unit", periodInput)
        else: 
            unit = atl2.replace("Unit!", unitinput)
        period = unit.replace("term!", periodInput)

        if studentinfo[self.col][7] == "M":
            return period.replace("!", "")
        else:
            p1 = period.replace("He!", "She")
            p2 = p1.replace("he!", "she")
            p3 = p2.replace("His!", "Her")
            p4 = p3.replace("his!", "her")
            return p4


stucount = len(studentinfo)
gradelist = []
totalMarks = []

if "load" not in st.session_state:
    st.session_state.load = True

def loadComments():
    if st.session_state.load:
        with st.spinner("Extending deadlines..."):
            bar = st.progress(0)
            time.sleep(0.3)
            bar.progress(30)
            time.sleep(0.5)
            bar.progress(40)
            time.sleep(0.2)
            bar.progress(50)
            time.sleep(0.3)
            for i in range(50, 80):
                time.sleep(0.03)
                bar.progress(i)
            bar.progress(100)
            #st.balloons()
            # Don"t do the animation a second time
            st.session_state.load = False

    for i in range(stucount):
        stx = student(i)
        gradelist.append(stx.finalGrade())
        totalMarks.append(stx.totalMarks())
        st.session_state[f"{stx.fn} {stx.ln} {collectiveInfo}"] = stx.finalComment()
        st.header(f"{stx.fn} {stx.ln} {collectiveInfo}")
        st.write(stx.finalComment())

loadComments()
# Visualization of Grades
markIndex = [i for i, x in enumerate(totalMarks) if x == max(totalMarks)]
counter = collections.Counter(gradelist)
cdict = dict(counter)

x = [1, 2, 3, 4, 5, 6, 7]
for i in x:
    if i not in cdict.keys():
        cdict[i] = 0

# dict = [
#     {"value": cdict[7], "name": "No. of 7"},
#     {"value": cdict[6], "name": "No. of 6"},
#     {"value": cdict[5], "name": "No. of 5"},
#     {"value": cdict[4], "name": "No. of 4"},
#     {"value": cdict[3], "name": "No. of 3"},
#     {"value": cdict[2], "name": "No. of 2"},
#     {"value": cdict[1], "name": "No. of 1"}
# ]
 
dict = [{"value": cdict[i], "name": f"No. of {i}"} for i in range(1,8)]


def hi():
    options = {
        "tooltip": {"trigger": "item"},
        "legend": {"top": "5%", "left": "center"},
        "series": [
            {
                "type": "pie",
                "radius": ["30%", "70%"],
                "avoidLabelOverlap": True,
                "itemStyle": {
                    "borderRadius": 10,
                    "borderColor": "#fff",
                    "borderWidth": 2,
                },
                "label": {"show": False, "position": "center"},
                "emphasis": {
                    "label": {"show": True, "fontSize": "40", "fontWeight": "bold"}
                },
                "labelLine": {"show": False},
                "data": dict,

            }
        ],
    }
    st_echarts(
        options=options, height="500px",
    )


st.header("Class Statistics")
with st.expander("Open statistics"):
    col1, col2 = st.columns(2)
    col1.metric(label="Class Size", value=stucount)
    col2.markdown("Top Students")
    for i in markIndex:
        col2.write(student(i).fn + " " + student(i).ln)

    hi()
